#!/usr/bin/env python3
"""Export src/data/*.json into a KDP-ready Word manuscript.

Usage: python3 scripts/generate_manuscript.py
Output: book/USCivicsPass-Workbook-Manuscript.docx
Requires: pip3 install python-docx
"""
import json
import os
import random
import re
import string
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DATA = os.path.join(ROOT, "src", "data")
OUT_DIR = os.path.join(ROOT, "book")
OUT_FILE = os.path.join(OUT_DIR, "USCivicsPass-Workbook-Manuscript.docx")


def load(name):
    with open(os.path.join(DATA, name), encoding="utf-8") as f:
        return json.load(f)


def set_page_setup(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)   # slightly wider for print binding gutter
    section.right_margin = Inches(0.75)


def new_section_page(doc):
    doc.add_section(WD_SECTION.NEW_PAGE)


def add_title_page(doc):
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_heading("New Complete U.S. Citizenship Test Prep Workbook", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph("Essential Study Guide for Success")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].italic = True
    subtitle.runs[0].font.size = Pt(16)

    doc.add_paragraph()
    features = [
        "100 & 128 Civics Practice Tests (with Answer Keys)",
        "N-400 Interview Guide (Process & Questions)",
        "Reading & Writing Practice (Tests & Key)",
        "Vocabulary & Word Bank (Definitions, Puzzles & Activities)",
        "Full Simulated Exams & Mock Interview Practice",
        "Cut-Out Flashcards for All 100 & 128 Questions",
    ]
    for line in features:
        p = doc.add_paragraph(f"• {line}")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    byline = doc.add_paragraph("By Hong Nguyen")
    byline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    byline.runs[0].bold = True
    byline.runs[0].font.size = Pt(14)


def add_copyright_page(doc):
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph("Copyright © 2026 by Hong Nguyen")
    p.runs[0].bold = True
    doc.add_paragraph("All rights reserved.")
    doc.add_paragraph(
        "No part of this publication may be reproduced, distributed, or transmitted "
        "in any form or by any means, including photocopying, recording, or other "
        "electronic or mechanical methods, without the prior written permission of "
        "the author, except in the case of brief quotations embodied in critical "
        "reviews and certain other noncommercial uses permitted by copyright law."
    )
    doc.add_paragraph()
    disclaimer_heading = doc.add_paragraph()
    disclaimer_heading.add_run("Disclaimer").bold = True
    doc.add_paragraph(
        "This workbook is an independent study guide created for educational "
        "purposes. It is not affiliated with, endorsed by, or sponsored by U.S. "
        "Citizenship and Immigration Services (USCIS) or any other government "
        "agency. All civics questions and answers are based on publicly available "
        "USCIS study materials current as of the publication date. Immigration "
        "laws, forms, fees, and procedures change frequently — always verify "
        "current requirements at the official USCIS website (uscis.gov) before "
        "applying or attending your interview."
    )
    doc.add_paragraph()
    doc.add_paragraph("First Edition, 2026")
    doc.add_paragraph("Printed in the United States of America")


def add_dedication_page(doc):
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    heading = doc.add_paragraph("Dedication")
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.runs[0].bold = True
    heading.runs[0].font.size = Pt(14)

    doc.add_paragraph()
    ded = doc.add_paragraph(
        "To everyone on the journey to becoming a U.S. citizen — your "
        "dedication, patience, and hope inspired this book. This workbook is "
        "for every future citizen studying late at night, practicing English "
        "out loud, and believing in a better future. You've got this."
    )
    ded.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ded.runs[0].italic = True

    doc.add_paragraph()
    doc.add_paragraph()
    note_heading = doc.add_paragraph("A Note from the Author")
    note_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note_heading.runs[0].bold = True
    note_heading.runs[0].font.size = Pt(14)

    doc.add_paragraph()
    note = doc.add_paragraph(
        "I created the USCivicsPass app and this companion workbook to make "
        "studying for the naturalization test less overwhelming. Every question, "
        "sentence, and word in this book was chosen to mirror the real interview "
        "experience as closely as possible, so there are no surprises on test "
        "day. I hope it helps you the way I wished a resource like this had "
        "helped people I know on their own citizenship journeys."
    )
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_how_to_use_section(doc):
    doc.add_heading("How to Use This Workbook", level=1)

    doc.add_heading("Test Overview", level=2)
    doc.add_paragraph(
        "USCIS naturalization applicants are tested on English (reading, "
        "writing, and speaking) and civics (U.S. history and government). "
        "Depending on when you filed and your interview notice, you may be "
        "tested on the 100-question (2008) version or the 128-question (2020) "
        "version of the civics test — this workbook includes complete "
        "practice material for both. Always confirm which version applies to "
        "you from your official USCIS interview notice or uscis.gov, since "
        "test versions have changed over time."
    )

    doc.add_heading("Suggested 8-Week Study Plan", level=2)
    plan = [
        ("Week 1", "Read the N-400 Interview Prep questions (Part 1). Focus on understanding, not memorizing."),
        ("Week 2", "Read the Mock Interview Script (Part 2) out loud."),
        ("Weeks 3–4", "Work through Reading Practice and Writing Practice (Parts 3 & 4) a little each day."),
        ("Week 5", "Study the Vocabulary & Word Bank (Part 5), then reinforce it with the puzzles and activities in Part 9."),
        ("Weeks 6–7", "Take 2–3 timed Practice Tests (Parts 6 & 7) under real interview conditions."),
        ("Week 8", "Take the full simulated exams and mock interview in Part 8, then review your Error Analysis Sheets."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "When"
    hdr[1].text = "What to Do"
    for week, task in plan:
        row = table.add_row().cells
        row[0].text = week
        row[1].text = task
    doc.add_paragraph()

    doc.add_heading("How This Book Is Organized", level=2)
    doc.add_paragraph(
        "Part 1–2 prepare you for the N-400 interview itself, including a "
        "realistic mock script. Part 3–4 build your reading and writing "
        "skills with real test sentences. Part 5 is a full vocabulary "
        "glossary. Parts 6–7 give you 20 complete multiple-choice practice "
        "tests with answer keys covering the 100- and 128-question civics "
        "test sets. Part 8 combines everything into full simulated exams "
        "and a fill-in-the-blank mock interview, with Error Analysis "
        "Sheets to track your progress. Part 9 reinforces the vocabulary "
        "glossary with matching and fill-in-the-blank puzzle activities. "
        "Part 10 gives you cut-out, fold-style flashcards for every "
        "question in both the 100- and 128-question sets."
    )


def add_roadmap_section(doc):
    doc.add_heading("U.S. Citizenship Journey Roadmap", level=1)
    doc.add_paragraph(
        "A general overview of the naturalization process. Timelines vary by "
        "case — always check your own USCIS case status for exact steps and dates."
    ).runs[0].italic = True
    doc.add_paragraph()

    steps = [
        ("1", "Check Your Eligibility", "Confirm you meet residency, age, and moral character requirements (see Quick Reference checklist)."),
        ("2", "File Form N-400", "Submit your Application for Naturalization online or by mail, with required documents and fee."),
        ("3", "Attend Biometrics", "USCIS schedules a fingerprinting and photo appointment for a background check."),
        ("4", "Study & Prepare", "Use this workbook to prepare for the civics, reading, and writing tests, and the interview."),
        ("5", "Attend Your Interview", "A USCIS officer reviews your application and administers the English and civics tests."),
        ("6", "Receive a Decision", "USCIS approves, continues, or denies your application, usually the same day or by mail."),
        ("7", "Take the Oath of Allegiance", "Attend your Oath ceremony — the final step to becoming a U.S. citizen."),
        ("8", "Become a U.S. Citizen!", "Receive your Certificate of Naturalization. Congratulations!"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Step"
    hdr[1].text = "Milestone"
    hdr[2].text = "What Happens"
    table.columns[0].width = Inches(0.5)
    for num, milestone, detail in steps:
        row = table.add_row().cells
        row[0].text = num
        row[1].text = milestone
        row[2].text = detail


def add_quick_reference_section(doc):
    doc.add_heading("Quick Reference", level=1)

    doc.add_heading("Eligibility Checklist", level=2)
    checklist = [
        "At least 18 years old",
        "Permanent resident (green card holder) for at least 5 years (or 3 years if married to a U.S. citizen)",
        "Continuous residence and physical presence in the U.S. for the required time",
        "Lived in your state or USCIS district for at least 3 months",
        "Good moral character",
        "Basic knowledge of U.S. government and history (civics test)",
        "Ability to read, write, and speak basic English",
        "Willingness to take the Oath of Allegiance",
    ]
    for item in checklist:
        doc.add_paragraph(f"☐  {item}")
    note = doc.add_paragraph(
        "Eligibility requirements vary by situation — confirm your specific "
        "case at uscis.gov/citizenship."
    )
    note.runs[0].italic = True

    doc.add_heading("Key Contacts & Resources", level=2)
    contacts = [
        ("USCIS Website", "uscis.gov"),
        ("USCIS Contact Center", "1-800-375-5283"),
        ("Case Status Online", "egov.uscis.gov/casestatus"),
        ("Find a Field Office", "uscis.gov/about-us/find-a-uscis-office"),
        ("Civics Test Study Materials", "uscis.gov/citizenship"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Resource"
    hdr[1].text = "Contact / Link"
    for name, value in contacts:
        row = table.add_row().cells
        row[0].text = name
        row[1].text = value
    doc.add_paragraph()

    doc.add_heading("My Key Dates", level=2)
    dates = [
        "N-400 Filing Date",
        "Biometrics Appointment",
        "Interview Date",
        "Oath Ceremony Date",
    ]
    for label in dates:
        p = doc.add_paragraph()
        p.add_run(f"{label}: ").bold = True
        p.add_run("_" * 30)


def add_toc_page(doc):
    doc.add_heading("Table of Contents", level=1)
    entries = [
        "Copyright & Disclaimer",
        "Dedication / Author's Note",
        "How to Use This Workbook",
        "U.S. Citizenship Journey Roadmap",
        "Quick Reference",
        "Part 1: N-400 Interview Prep",
        "Part 2: Mock Interview Script",
        "Part 3: Reading Practice",
        "Part 4: Writing Practice",
        "Part 5: Vocabulary & Word Bank",
        "Part 6: 100-Question Practice Tests (Sets 1-10)",
        "Part 7: 128-Question Practice Tests (Sets 1-10)",
        "Part 8: Full Practice Tests & Mock Interviews",
        "Part 9: Vocabulary Puzzles & Activities",
        "Part 10: Cut-Out Flashcards (100 & 128 Questions)",
        "Back Matter",
    ]
    for e in entries:
        doc.add_paragraph(e, style="List Number")
    sub_entries = [
        "Mock Interview Simulation (fill-in-the-blank) (Part 8)",
        "10 Complete Practice Tests (100 Civics + Reading/Writing) (Part 8)",
        "10 Complete Practice Tests (128 Civics + Reading/Writing) (Part 8)",
        "Review & Error Analysis Sheets (Part 8)",
        "Matching Activities, Fill-in-the-Blank (Part 9)",
        "100-Question Flashcards (Part 10)",
        "128-Question Flashcards (Part 10)",
        "Glossary of Key Terms (Back Matter)",
        "Motivation (Back Matter)",
        "Notes Pages (Back Matter)",
        "Progress Certificate (Back Matter)",
        "A Final Word from the Author (Back Matter)",
    ]
    for e in sub_entries:
        p = doc.add_paragraph(e, style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.5)


def add_n400_section(doc, data):
    doc.add_heading("Part 1: N-400 Interview Prep", level=1)
    current_category = None
    for item in data:
        if item.get("category") != current_category:
            current_category = item.get("category")
            doc.add_heading(current_category, level=2)
        p = doc.add_paragraph()
        run = p.add_run(f"{item['id']}. {item['question']}")
        run.bold = True
        ans = doc.add_paragraph()
        ans.paragraph_format.left_indent = Inches(0.25)
        ans_run = ans.add_run(f"Answer: {item['answer']}")
        ans_run.italic = True
        doc.add_paragraph()


def add_mock_interview_section(doc, data):
    doc.add_heading("Part 2: Mock Interview Script", level=1)
    for item in data:
        doc.add_heading(item["section"], level=2)
        for line in item["text"].split("\n"):
            line = line.strip()
            if not line:
                continue
            p = doc.add_paragraph()
            if ":" in line:
                speaker, rest = line.split(":", 1)
                run = p.add_run(f"{speaker}:")
                run.bold = True
                p.add_run(rest)
            else:
                p.add_run(line)
        doc.add_paragraph()


def add_sentence_practice_section(doc, title, sentences, with_writing_lines=False):
    doc.add_heading(title, level=1)
    for i, sentence in enumerate(sentences, start=1):
        p = doc.add_paragraph()
        p.add_run(f"{i}. {sentence}")
        if with_writing_lines:
            blank = doc.add_paragraph()
            blank.paragraph_format.left_indent = Inches(0.25)
            blank.add_run("_" * 70)
        doc.add_paragraph()


def build_practice_sets(pool, questions_per_set, num_sets, seed):
    """Split pool into num_sets practice tests of questions_per_set each.

    If the sets partition the pool exactly (e.g. 10 sets of 10 from a
    100-question pool), each question appears exactly once across all sets.
    Otherwise (e.g. 10 sets of 20 from a 128-question pool) sets are drawn
    independently, so some overlap across sets is expected -- this mirrors
    how real prep books reuse questions across practice tests.
    """
    rng = random.Random(seed)
    total_needed = questions_per_set * num_sets
    if total_needed == len(pool):
        shuffled = pool[:]
        rng.shuffle(shuffled)
        return [shuffled[i * questions_per_set:(i + 1) * questions_per_set] for i in range(num_sets)]
    return [rng.sample(pool, questions_per_set) for _ in range(num_sets)]


LETTERS = ["A", "B", "C"]


def generate_options(rng, correct_answer, pool):
    """Mirror App.jsx's generateOptions: 1 correct + 2 random wrong answers
    pulled from other questions in the same pool, shuffled together."""
    wrong_answers = [q["answer"] for q in pool if q["answer"] != correct_answer]
    rng.shuffle(wrong_answers)
    wrong_answers = wrong_answers[:2]
    options = [correct_answer] + wrong_answers
    rng.shuffle(options)
    return options


def add_practice_tests_section(doc, title, pool, questions_per_set, num_sets, seed):
    doc.add_heading(title, level=1)
    intro = doc.add_paragraph(
        f"Each practice test below has {questions_per_set} multiple-choice questions, "
        f"just like the USCivicsPass app quiz. Circle your answer, then check the "
        f"Answer Key to see how you did."
    )
    intro.runs[0].italic = True
    doc.add_paragraph()

    sets = build_practice_sets(pool, questions_per_set, num_sets, seed)
    for set_idx, qset in enumerate(sets, start=1):
        rng = random.Random(f"{seed}-{set_idx}")
        doc.add_heading(f"Practice Test {set_idx}", level=2)

        correct_letters = []
        for i, item in enumerate(qset, start=1):
            options = generate_options(rng, item["answer"], pool)
            correct_letters.append(LETTERS[options.index(item["answer"])])

            p = doc.add_paragraph()
            run = p.add_run(f"{i}. {item['question']}")
            run.bold = True
            for letter, option in zip(LETTERS, options):
                opt_p = doc.add_paragraph()
                opt_p.paragraph_format.left_indent = Inches(0.35)
                opt_p.add_run(f"{letter}) {option}")

        doc.add_paragraph()
        doc.add_heading("Answer Key", level=3)
        key_p = doc.add_paragraph()
        key_p.add_run(
            "   ".join(f"{i}. {letter}" for i, letter in enumerate(correct_letters, start=1))
        )

        if set_idx != num_sets:
            doc.add_page_break()


def add_mock_interview_simulation_section(doc, data):
    doc.add_heading("Mock Interview Simulation", level=2)
    intro = doc.add_paragraph(
        "Practice out loud. Read each Officer line, then say your answer before "
        "writing it on the blank line below — just like the real interview."
    )
    intro.runs[0].italic = True
    doc.add_paragraph()

    for item in data:
        doc.add_heading(item["section"], level=3)
        for line in item["text"].split("\n"):
            line = line.strip()
            if not line:
                continue
            speaker, _, rest = line.partition(":")
            if speaker.strip().lower().startswith("officer"):
                p = doc.add_paragraph()
                p.add_run(f"{speaker}:").bold = True
                p.add_run(rest)
            elif speaker.strip().lower().startswith("applicant"):
                label = doc.add_paragraph()
                label.add_run("Applicant:").bold = True
                blank = doc.add_paragraph()
                blank.paragraph_format.left_indent = Inches(0.25)
                blank.add_run("_" * 70)
            else:
                p = doc.add_paragraph()
                p.add_run(line)
        doc.add_paragraph()


def sample_sentence_sets(pool, count, num_sets, seed):
    rng = random.Random(seed)
    total_needed = count * num_sets
    if total_needed == len(pool):
        shuffled = pool[:]
        rng.shuffle(shuffled)
        return [shuffled[i * count:(i + 1) * count] for i in range(num_sets)]
    return [rng.sample(pool, count) for _ in range(num_sets)]


def add_complete_practice_tests_section(doc, title, civics_pool, questions_per_set,
                                         reading_pool, writing_pool, num_tests, seed):
    doc.add_heading(title, level=2)
    intro = doc.add_paragraph(
        f"Each complete practice test below simulates the real naturalization "
        f"interview: {questions_per_set} civics questions, a reading test, and a "
        f"writing test."
    )
    intro.runs[0].italic = True
    doc.add_paragraph()

    civics_sets = build_practice_sets(civics_pool, questions_per_set, num_tests, f"{seed}-civics")
    reading_sets = sample_sentence_sets(reading_pool, 3, num_tests, f"{seed}-reading")
    writing_sets = sample_sentence_sets(writing_pool, 3, num_tests, f"{seed}-writing")

    for test_idx in range(1, num_tests + 1):
        rng = random.Random(f"{seed}-{test_idx}")
        doc.add_heading(f"Complete Practice Test {test_idx}", level=3)

        doc.add_heading("Civics Test", level=4)
        qset = civics_sets[test_idx - 1]
        correct_letters = []
        for i, item in enumerate(qset, start=1):
            options = generate_options(rng, item["answer"], civics_pool)
            correct_letters.append(LETTERS[options.index(item["answer"])])
            p = doc.add_paragraph()
            run = p.add_run(f"{i}. {item['question']}")
            run.bold = True
            for letter, option in zip(LETTERS, options):
                opt_p = doc.add_paragraph()
                opt_p.paragraph_format.left_indent = Inches(0.35)
                opt_p.add_run(f"{letter}) {option}")

        doc.add_heading("Reading Test", level=4)
        doc.add_paragraph("Read each sentence aloud to a partner or record yourself:")
        for i, sentence in enumerate(reading_sets[test_idx - 1], start=1):
            doc.add_paragraph(f"{i}. {sentence}")

        doc.add_heading("Writing Test", level=4)
        doc.add_paragraph("Listen to (or read) each sentence, then write it below:")
        for i, sentence in enumerate(writing_sets[test_idx - 1], start=1):
            p = doc.add_paragraph()
            p.add_run(f"{i}. {sentence}")
            blank = doc.add_paragraph()
            blank.paragraph_format.left_indent = Inches(0.25)
            blank.add_run("_" * 70)

        doc.add_paragraph()
        doc.add_heading("Civics Answer Key", level=4)
        key_p = doc.add_paragraph()
        key_p.add_run(
            "   ".join(f"{i}. {letter}" for i, letter in enumerate(correct_letters, start=1))
        )

        if test_idx != num_tests:
            doc.add_page_break()


def add_review_error_analysis_section(doc, num_sheets):
    doc.add_heading("Review & Error Analysis Sheets", level=2)
    intro = doc.add_paragraph(
        "After each complete practice test, use a sheet below to log every "
        "question you missed. Reviewing your mistakes is the fastest way to "
        "improve your score."
    )
    intro.runs[0].italic = True
    doc.add_paragraph()

    for sheet_num in range(1, num_sheets + 1):
        doc.add_heading(f"Error Analysis Sheet {sheet_num}", level=3)
        p = doc.add_paragraph()
        p.add_run("Practice Test #: ").bold = True
        p.add_run("_" * 10)
        p.add_run("     Score: ").bold = True
        p.add_run("_" * 10)

        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "Q#"
        hdr[1].text = "My Answer"
        hdr[2].text = "Correct Answer"
        hdr[3].text = "Why I Missed It"
        for _ in range(12):
            table.add_row()

        doc.add_paragraph()
        notes = doc.add_paragraph()
        notes.add_run("Topics to review before my next practice test:").bold = True
        for _ in range(3):
            blank = doc.add_paragraph()
            blank.add_run("_" * 80)

        if sheet_num != num_sheets:
            doc.add_page_break()


def add_vocabulary_section(doc, data):
    doc.add_heading("Part 5: Vocabulary & Word Bank", level=1)
    sorted_words = sorted(data, key=lambda x: x["word"].lower())
    current_letter = None
    for entry in sorted_words:
        letter = entry["word"][0].upper()
        if letter != current_letter:
            current_letter = letter
            doc.add_heading(current_letter, level=2)
        p = doc.add_paragraph()
        run = p.add_run(entry["word"])
        run.bold = True
        p.add_run(f" — {entry['meaning']}")


# ---------------------------------------------------------------------------
# Part 9: Vocabulary Puzzles & Activities
# ---------------------------------------------------------------------------

class WordQueue:
    """Draws n items at a time from pool, guaranteeing every item is used
    once before any item repeats (reshuffles and refills on exhaustion)."""

    def __init__(self, pool, seed):
        self.pool = list(pool)
        self.rng = random.Random(seed)
        self.queue = []

    def draw(self, n):
        while len(self.queue) < n:
            batch = list(self.pool)
            self.rng.shuffle(batch)
            self.queue.extend(batch)
        result, self.queue = self.queue[:n], self.queue[n:]
        return result


def normalize_word(raw):
    """Vocabulary entries include alternatives ('A / B') and parentheticals.
    Take the first/main phrase, and derive a letters-only grid token."""
    display = raw.split("/")[0].split("(")[0].strip()
    token = re.sub(r"[^A-Za-z]", "", display).upper()
    return display, token


def add_worksheet_header(doc, title):
    """Name/Date line + a large centered worksheet title, matching the
    Super Teacher Worksheets-style layout used across the activity pages."""
    nd = doc.add_paragraph()
    nd.add_run("Name: ").bold = True
    nd.add_run("_" * 30 + "    ")
    nd.add_run("Date: ").bold = True
    nd.add_run("_" * 12)

    heading = doc.add_paragraph(title)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.runs[0].bold = True
    heading.runs[0].font.size = Pt(24)
    doc.add_paragraph()


def add_matching_subsection(doc, vocab_pool, seed, num_pages=10, per_page=25):
    doc.add_heading("Matching Activities", level=2)
    doc.add_paragraph(
        "Match each word to its correct definition. Write the letter of the "
        "definition on the blank next to each word."
    ).runs[0].italic = True
    doc.add_paragraph()

    queue = WordQueue(vocab_pool, seed)
    letters = [chr(ord("A") + k) for k in range(per_page)]
    cache = []
    for i in range(1, num_pages + 1):
        entries = queue.draw(per_page)
        rng = random.Random(f"{seed}-{i}")
        def_order = list(range(per_page))
        rng.shuffle(def_order)
        cache.append((i, entries, def_order))

        doc.add_heading(f"Matching Activity {i}", level=3)
        add_worksheet_header(doc, "Civics Vocabulary")
        instr = doc.add_paragraph()
        instr.add_run(
            "Match each vocabulary term to its definition by writing the "
            "correct letter on the line."
        ).italic = True
        doc.add_paragraph()

        table = doc.add_table(rows=per_page + 1, cols=2)
        table.style = "Table Grid"
        hdr_word, hdr_def = table.rows[0].cells
        hdr_word.paragraphs[0].add_run("Vocabulary Term").bold = True
        hdr_def.paragraphs[0].add_run("Definitions").bold = True
        for row_idx in range(per_page):
            word_display = entries[row_idx][0]
            table.cell(row_idx + 1, 0).text = f"{row_idx + 1}. {word_display}  ____"
            meaning = entries[def_order[row_idx]][2]
            table.cell(row_idx + 1, 1).text = f"{letters[row_idx]}. {meaning}"
        if i != num_pages:
            doc.add_page_break()

    doc.add_heading("Matching Answer Keys", level=3)
    for i, entries, def_order in cache:
        inverse = [0] * per_page
        for right_row, word_idx in enumerate(def_order):
            inverse[word_idx] = right_row
        doc.add_heading(f"Matching Activity {i}", level=4)
        key_p = doc.add_paragraph()
        key_p.add_run(
            "   ".join(f"{n}. {letters[inverse[n - 1]]}" for n in range(1, per_page + 1))
        )


def render_boxed_word_bank(doc, words, columns=4):
    rows = 1 + (len(words) + columns - 1) // columns
    table = doc.add_table(rows=rows, cols=columns)
    table.style = "Table Grid"
    header_cell = table.cell(0, 0)
    for c in range(1, columns):
        header_cell = header_cell.merge(table.cell(0, c))
    header_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header_cell.paragraphs[0].add_run("Word Bank")
    header_run.bold = True
    header_run.font.size = Pt(13)
    for i, word in enumerate(words):
        r, c = divmod(i, columns)
        cell = table.cell(r + 1, c)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(word)
    return table


def add_fill_in_blank_subsection(doc, vocab_pool, seed, num_pages=8, per_page=25):
    doc.add_heading("Fill-in-the-Blank Activities", level=2)
    doc.add_paragraph(
        "Use the Word Bank to fill in the best answer for each question below."
    ).runs[0].italic = True
    doc.add_paragraph()

    queue = WordQueue(vocab_pool, seed)
    cache = []
    for i in range(1, num_pages + 1):
        entries = queue.draw(per_page)
        rng = random.Random(f"{seed}-{i}")
        order = list(range(per_page))
        rng.shuffle(order)
        ordered_entries = [entries[k] for k in order]
        cache.append((i, ordered_entries))

        doc.add_heading(f"Fill-in-the-Blank {i}", level=3)
        add_worksheet_header(doc, "U.S. Civics")
        instr = doc.add_paragraph()
        instr.add_run("Fill in the best answer for each question.").italic = True
        doc.add_paragraph()

        bank_words = sorted((e[0] for e in entries), key=str.lower)
        render_boxed_word_bank(doc, bank_words, columns=4)
        doc.add_paragraph()

        for idx, (display, token, meaning) in enumerate(ordered_entries, start=1):
            m = meaning[0].lower() + meaning[1:] if meaning else meaning
            doc.add_paragraph(f"{idx}. The term for \"{m}\" is {'_' * 20}.")
        if i != num_pages:
            doc.add_page_break()

    doc.add_heading("Fill-in-the-Blank Answer Keys", level=3)
    for i, ordered_entries in cache:
        doc.add_heading(f"Fill-in-the-Blank {i}", level=4)
        for idx, (display, token, meaning) in enumerate(ordered_entries, start=1):
            doc.add_paragraph(f"{idx}. {display}")


def add_vocabulary_activities_section(doc, vocab):
    doc.add_heading("Part 9: Vocabulary Puzzles & Activities", level=1)
    doc.add_paragraph(
        "254 vocabulary words, reinforced through two activity types so you "
        "encounter each word multiple times in different formats: matching "
        "and fill-in-the-blank. Each activity type has its own answer key at "
        "the end of that section."
    ).runs[0].italic = True
    doc.add_page_break()

    matching_pool = [normalize_word(e["word"]) + (e["meaning"],) for e in vocab]
    add_matching_subsection(doc, matching_pool, seed="match")
    doc.add_page_break()

    add_fill_in_blank_subsection(doc, matching_pool, seed="fib")


def add_bottom_border(paragraph):
    """Give a paragraph a bottom rule, used to render blank lined pages."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    pBdr.append(bottom)
    pPr.append(pBdr)


def set_cell_border(cell, edges=("top", "bottom", "left", "right"), val="dashed", sz="8", color="595959"):
    """Apply a border (default: dashed, used as a cut line) to a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge in edges:
        tag = f"w:{edge}"
        el = tcBorders.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            tcBorders.append(el)
        el.set(qn("w:val"), val)
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def add_dashed_bottom_border(paragraph, color="595959"):
    """Give a paragraph a dashed bottom rule, used as the fold line on a flashcard."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "dashed")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def set_cell_margins(cell, top=100, bottom=100, left=120, right=120):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, value in (("top", top), ("bottom", bottom), ("start", left), ("end", right)):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)


def add_flashcard_cell(cell, item):
    set_cell_border(cell)
    set_cell_margins(cell)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    q_p = cell.paragraphs[0]
    q_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    q_run = q_p.add_run(f"Q{item['id']}. {item['question']}")
    q_run.bold = True
    q_run.font.size = Pt(10)

    fold_p = cell.add_paragraph()
    fold_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_dashed_bottom_border(fold_p)
    fold_run = fold_p.add_run("✂ - - - - - -  fold here  - - - - - - ✂")
    fold_run.italic = True
    fold_run.font.size = Pt(7)
    fold_run.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

    a_p = cell.add_paragraph()
    a_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    a_run = a_p.add_run(f"A: {item['answer']}")
    a_run.italic = True
    a_run.font.size = Pt(9)


def add_flashcards_section(doc, title, pool, cols=2, rows_per_page=3):
    """Render a pool of Q&A items as a grid of cut-out, fold-style flashcards.

    Each card has a dashed cut-line border and a dashed fold line between the
    question and answer, so a single card works as a flashcard: cut along the
    border, fold along the middle, and the answer is hidden on the back until
    you flip it up.

    KDP interiors are printed double-sided, so every page has *something*
    printed on its physical back by default. To keep that back blank (so a
    cut-out card isn't printed on the reverse), each flashcard grid page is
    forced onto an odd (recto) page via an ODD_PAGE section break — Word/KDP
    automatically inserts a blank even (verso) page whenever needed to make
    that happen, which becomes the physical back of the previous sheet.
    """
    doc.add_heading(title, level=2)
    intro = doc.add_paragraph(
        "Cut out each card along the dashed border. Fold along the dashed "
        "line in the middle so the question is on the front and the answer "
        "is tucked underneath — flip it up to check yourself. Each flashcard "
        "page is followed by a blank page, so the back of every printed "
        "sheet stays empty once you cut the cards apart."
    )
    intro.runs[0].italic = True
    doc.add_paragraph()

    per_page = cols * rows_per_page
    col_width = Inches(3.35)
    total = len(pool)
    chunks = [pool[i:i + per_page] for i in range(0, total, per_page)]

    for chunk in chunks:
        # Force this card page onto the next odd page; if the previous page
        # was odd, this auto-inserts one blank even page as its back.
        doc.add_section(WD_SECTION.ODD_PAGE)
        table = doc.add_table(rows=0, cols=cols)
        table.style = None
        table.autofit = False
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for i in range(0, len(chunk), cols):
            row_items = chunk[i:i + cols]
            row = table.add_row()
            row.height = Inches(2.6)
            row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            for col_idx in range(cols):
                cell = row.cells[col_idx]
                cell.width = col_width
                if col_idx < len(row_items):
                    add_flashcard_cell(cell, row_items[col_idx])

    # Push whatever comes next to the following odd page too, so the last
    # flashcard sheet's back is also left blank.
    doc.add_section(WD_SECTION.ODD_PAGE)


def add_glossary_section(doc, vocab):
    doc.add_heading("Glossary of Key Terms", level=1)
    intro = doc.add_paragraph(
        "A quick-reference glossary of every civics vocabulary term in this "
        "workbook, grouped alphabetically for easy lookup after you've "
        "finished Part 5."
    )
    intro.runs[0].italic = True
    doc.add_paragraph()

    sorted_words = sorted(vocab, key=lambda x: x["word"].lower())
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Term"
    hdr[1].text = "Definition"
    table.columns[0].width = Inches(2.0)

    current_letter = None
    for entry in sorted_words:
        letter = entry["word"][0].upper()
        if letter != current_letter:
            current_letter = letter
            letter_row = table.add_row().cells
            merged = letter_row[0].merge(letter_row[1])
            merged.paragraphs[0].add_run(current_letter).bold = True
        row = table.add_row().cells
        row[0].paragraphs[0].add_run(entry["word"]).bold = True
        row[1].paragraphs[0].add_run(entry["meaning"])


def add_motivation_section(doc):
    doc.add_heading("Motivation", level=1)
    heading = doc.add_paragraph("You've Made It This Far")
    heading.runs[0].bold = True
    heading.runs[0].font.size = Pt(16)
    doc.add_paragraph()

    body = [
        "If you're reading this page, you've already worked through civics "
        "questions, reading and writing practice, vocabulary, and full "
        "practice tests. That's not a small thing — it's hours of focus on "
        "top of work, family, and everything else in your life.",
        "The naturalization interview can feel intimidating, but "
        "preparation is what turns nerves into confidence. You already know "
        "more than you think you do. Trust the work you've put in.",
    ]
    for para in body:
        doc.add_paragraph(para)

    doc.add_paragraph()
    prompt = doc.add_paragraph()
    prompt.add_run("Why I want to become a U.S. citizen:").bold = True
    for _ in range(6):
        line = doc.add_paragraph()
        line.paragraph_format.space_after = Pt(18)
        add_bottom_border(line)


def add_notes_section(doc, num_pages=4, lines_per_page=26):
    doc.add_heading("Notes", level=1)
    intro = doc.add_paragraph(
        "Use these pages for questions to ask your legal representative, "
        "reminders for interview day, or anything else you want to remember."
    )
    intro.runs[0].italic = True
    doc.add_paragraph()

    for page in range(1, num_pages + 1):
        for _ in range(lines_per_page):
            line = doc.add_paragraph()
            line.paragraph_format.space_after = Pt(18)
            add_bottom_border(line)
        if page != num_pages:
            doc.add_page_break()


def add_progress_certificate_section(doc):
    doc.add_heading("Progress Certificate", level=1)
    intro = doc.add_paragraph(
        "Fill this out once you've completed the full workbook — a "
        "milestone worth celebrating before your interview."
    )
    intro.runs[0].italic = True
    doc.add_paragraph()
    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)

    p0 = cell.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run0 = p0.add_run("Certificate of Completion")
    run0.bold = True
    run0.font.size = Pt(24)

    cell.add_paragraph()

    p1 = cell.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.add_run("This certifies that").italic = True

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run("_" * 40)

    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.add_run(
        "has successfully completed the New Complete U.S. Citizenship Test "
        "Prep Workbook and is ready for the naturalization interview."
    ).italic = True

    cell.add_paragraph()

    p4 = cell.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.add_run("Date: ").bold = True
    p4.add_run("_" * 20 + "     ")
    p4.add_run("Signature: ").bold = True
    p4.add_run("_" * 20)

    cell.add_paragraph()


def add_hyperlink(paragraph, url, text, bold=False):
    """Insert a real clickable hyperlink run into a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1155CC")
    rPr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rPr.append(underline)
    if bold:
        rPr.append(OxmlElement("w:b"))
    run.append(rPr)

    text_el = OxmlElement("w:t")
    text_el.text = text
    run.append(text_el)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


def add_final_words_section(doc):
    doc.add_heading("A Final Word from the Author", level=1)
    doc.add_paragraph()

    body = [
        "Thank you for choosing this workbook as part of your journey to "
        "U.S. citizenship. Whether you worked through every page or focused "
        "on just the sections you needed most, I hope it helped you walk "
        "into your interview with confidence.",
        "This workbook is the companion to the USCivicsPass app — a free "
        "study tool with interactive quizzes, flashcards, audio "
        "pronunciation, and mock interview practice covering the same "
        "material in this book. You can keep practicing anytime, on any "
        "device, at:",
    ]
    for para in body:
        doc.add_paragraph(para)

    link_p = doc.add_paragraph()
    link_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_hyperlink(
        link_p, "https://uscivicspass.netlify.app/",
        "https://uscivicspass.netlify.app/", bold=True,
    )

    doc.add_paragraph()
    doc.add_paragraph(
        "Wishing you every success on your interview and on the rest of "
        "your journey. Congratulations in advance — you've earned it."
    )

    doc.add_paragraph()
    sign = doc.add_paragraph("— Hong Nguyen")
    sign.runs[0].bold = True
    sign.runs[0].italic = True


def main():
    os.makedirs(OUT_DIR, exist_ok=True)

    q100 = load("questions.json")
    q128 = load("questions128.json")
    n400 = load("n400Questions.json")
    mock_interview = load("mockInterview.json")
    reading = load("readingSentences.json")
    writing = load("writingSentences.json")
    vocab = load("citizenshipVocabulary.json")

    doc = Document()
    set_page_setup(doc)

    add_title_page(doc)
    doc.add_page_break()

    add_copyright_page(doc)
    doc.add_page_break()

    add_dedication_page(doc)
    doc.add_page_break()

    add_how_to_use_section(doc)
    doc.add_page_break()

    add_roadmap_section(doc)
    doc.add_page_break()

    add_quick_reference_section(doc)
    doc.add_page_break()

    add_toc_page(doc)
    doc.add_page_break()

    add_n400_section(doc, n400)
    doc.add_page_break()

    add_mock_interview_section(doc, mock_interview)
    doc.add_page_break()

    add_sentence_practice_section(doc, "Part 3: Reading Practice", reading, with_writing_lines=False)
    doc.add_page_break()

    add_sentence_practice_section(doc, "Part 4: Writing Practice", writing, with_writing_lines=True)
    doc.add_page_break()

    add_vocabulary_section(doc, vocab)
    doc.add_page_break()

    add_practice_tests_section(
        doc, "Part 6: 100-Question Practice Tests (Sets 1-10)",
        q100, questions_per_set=10, num_sets=10, seed=100,
    )
    doc.add_page_break()

    add_practice_tests_section(
        doc, "Part 7: 128-Question Practice Tests (Sets 1-10)",
        q128, questions_per_set=20, num_sets=10, seed=128,
    )
    doc.add_page_break()

    doc.add_heading("Part 8: Full Practice Tests & Mock Interviews", level=1)
    add_mock_interview_simulation_section(doc, mock_interview)
    doc.add_page_break()

    add_complete_practice_tests_section(
        doc, "10 Complete Practice Tests (100 Civics + Reading/Writing)",
        q100, questions_per_set=10, reading_pool=reading, writing_pool=writing,
        num_tests=10, seed="complete100",
    )
    doc.add_page_break()

    add_complete_practice_tests_section(
        doc, "10 Complete Practice Tests (128 Civics + Reading/Writing)",
        q128, questions_per_set=20, reading_pool=reading, writing_pool=writing,
        num_tests=10, seed="complete128",
    )
    doc.add_page_break()

    add_review_error_analysis_section(doc, num_sheets=20)
    doc.add_page_break()

    add_vocabulary_activities_section(doc, vocab)
    doc.add_page_break()

    doc.add_heading("Part 10: Cut-Out Flashcards", level=1)
    intro = doc.add_paragraph(
        "Flashcards for every question in both the 100- and 128-question "
        "civics sets, ready to cut out and carry with you."
    )
    intro.runs[0].italic = True
    doc.add_page_break()

    add_flashcards_section(doc, "100-Question Flashcards", q100)
    add_flashcards_section(doc, "128-Question Flashcards", q128)

    doc.add_heading("Back Matter", level=1)
    doc.add_page_break()

    add_glossary_section(doc, vocab)
    doc.add_page_break()

    add_motivation_section(doc)
    doc.add_page_break()

    add_notes_section(doc)
    doc.add_page_break()

    add_progress_certificate_section(doc)
    doc.add_page_break()

    add_final_words_section(doc)

    doc.save(OUT_FILE)
    print(f"Manuscript written to {OUT_FILE}")
    print(f"Sections: {len(q100)} + {len(q128)} civics questions, "
          f"{len(n400)} N-400 questions, {len(mock_interview)} interview scenes, "
          f"{len(reading)} reading sentences, {len(writing)} writing sentences, "
          f"{len(vocab)} vocabulary entries, "
          f"10 practice tests x 10Q (100-set) + 10 practice tests x 20Q (128-set).")


if __name__ == "__main__":
    main()
